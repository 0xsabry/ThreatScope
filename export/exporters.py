"""
ThreatScope V2 — Multi-Format Export Engine
Author: 0xSABRY

Generates reports in PDF, DOCX, STIX 2.1, JSON, and CSV formats.
"""

import json
import csv
import uuid
import logging
from pathlib import Path
from datetime import datetime, timezone
from typing import Dict, List, Optional

logger = logging.getLogger("threatscope.export")

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    )
    from reportlab.lib.units import inch, mm
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def export_json(results: dict, filepath: str) -> str:
    """
    Export analysis results as JSON.

    Args:
        results: Complete analysis results.
        filepath: Output file path.

    Returns:
        Path to exported file.
    """
    output = {
        "tool": "ThreatScope V2",
        "author": "0xSABRY",
        "export_time": datetime.now(timezone.utc).isoformat(),
        "results": _serialize_results(results),
    }

    with open(filepath, "w", encoding="utf-8") as f:
        json.dump(output, f, indent=2, default=str)
    logger.info(f"JSON report exported to {filepath}")
    return filepath


def export_csv(results: dict, filepath: str) -> str:
    """
    Export findings as CSV.

    Args:
        results: Complete analysis results.
        filepath: Output file path.

    Returns:
        Path to exported file.
    """
    findings = results.get("findings", [])
    if not findings:
        findings = [{"message": "No findings"}]

    fieldnames = ["severity", "title", "description", "mitre", "timestamp",
                  "line_number", "category_key", "rule_type", "raw"]

    with open(filepath, "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames, extrasaction="ignore")
        writer.writeheader()
        for finding in findings:
            writer.writerow(finding)

    logger.info(f"CSV report exported to {filepath}")
    return filepath


def export_stix(results: dict, filepath: str) -> str:
    """
    Export IOCs as STIX 2.1 bundle.

    Args:
        results: Complete analysis results.
        filepath: Output file path.

    Returns:
        Path to exported file.
    """
    stix_objects = []

    # Identity object
    identity_id = f"identity--{uuid.uuid5(uuid.NAMESPACE_DNS, 'threatscope')}"
    stix_objects.append({
        "type": "identity",
        "spec_version": "2.1",
        "id": identity_id,
        "created": datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ"),
        "modified": datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ"),
        "name": "ThreatScope V2",
        "identity_class": "tool",
    })

    iocs = results.get("iocs", {})
    now = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")

    # IP indicators
    for ip in iocs.get("top_ips", []):
        obj_id = f"indicator--{uuid.uuid5(uuid.NAMESPACE_DNS, f'ip:{ip}')}"
        stix_objects.append({
            "type": "indicator",
            "spec_version": "2.1",
            "id": obj_id,
            "created": now,
            "modified": now,
            "name": f"Suspicious IP: {ip}",
            "pattern": f"[ipv4-addr:value = '{ip}']",
            "pattern_type": "stix",
            "valid_from": now,
            "indicator_types": ["malicious-activity"],
            "created_by_ref": identity_id,
        })

    # Hash indicators
    for hash_type in ["md5", "sha1", "sha256"]:
        hashes = iocs.get("hashes", {}).get(hash_type, [])
        for h in hashes[:20]:
            obj_id = f"indicator--{uuid.uuid5(uuid.NAMESPACE_DNS, f'{hash_type}:{h}')}"
            stix_objects.append({
                "type": "indicator",
                "spec_version": "2.1",
                "id": obj_id,
                "created": now,
                "modified": now,
                "name": f"File Hash ({hash_type.upper()}): {h[:16]}...",
                "pattern": f"[file:hashes.'{hash_type.upper()}' = '{h}']",
                "pattern_type": "stix",
                "valid_from": now,
                "indicator_types": ["malicious-activity"],
                "created_by_ref": identity_id,
            })

    # Domain indicators
    for domain in iocs.get("top_domains", []):
        obj_id = f"indicator--{uuid.uuid5(uuid.NAMESPACE_DNS, f'domain:{domain}')}"
        stix_objects.append({
            "type": "indicator",
            "spec_version": "2.1",
            "id": obj_id,
            "created": now,
            "modified": now,
            "name": f"Suspicious Domain: {domain}",
            "pattern": f"[domain-name:value = '{domain}']",
            "pattern_type": "stix",
            "valid_from": now,
            "indicator_types": ["malicious-activity"],
            "created_by_ref": identity_id,
        })

    # URL indicators
    for url in iocs.get("urls", [])[:20]:
        obj_id = f"indicator--{uuid.uuid5(uuid.NAMESPACE_DNS, f'url:{url}')}"
        stix_objects.append({
            "type": "indicator",
            "spec_version": "2.1",
            "id": obj_id,
            "created": now,
            "modified": now,
            "name": f"Suspicious URL: {url[:50]}...",
            "pattern": f"[url:value = '{url}']",
            "pattern_type": "stix",
            "valid_from": now,
            "indicator_types": ["malicious-activity"],
            "created_by_ref": identity_id,
        })

    # Attack patterns (from MITRE hits)
    for technique, count in results.get("mitre_hits", {}).items():
        obj_id = f"attack-pattern--{uuid.uuid5(uuid.NAMESPACE_DNS, f'mitre:{technique}')}"
        stix_objects.append({
            "type": "attack-pattern",
            "spec_version": "2.1",
            "id": obj_id,
            "created": now,
            "modified": now,
            "name": f"MITRE ATT&CK: {technique}",
            "external_references": [{
                "source_name": "mitre-attack",
                "external_id": technique,
                "url": f"https://attack.mitre.org/techniques/{technique.replace('.', '/')}/",
            }],
        })

    # STIX Bundle
    bundle = {
        "type": "bundle",
        "id": f"bundle--{uuid.uuid4()}",
        "objects": stix_objects,
    }

    with open(filepath, "w", encoding="utf-8") as f:
        json.dump(bundle, f, indent=2)

    logger.info(f"STIX 2.1 bundle exported to {filepath} ({len(stix_objects)} objects)")
    return filepath


def export_pdf(results: dict, filepath: str) -> str:
    """
    Export professional PDF report with executive and technical sections.

    Args:
        results: Complete analysis results.
        filepath: Output file path.

    Returns:
        Path to exported file.
    """
    if not PDF_AVAILABLE:
        logger.warning("reportlab not installed — PDF export unavailable")
        # Fallback to text-based report
        txt_path = filepath.replace(".pdf", ".txt")
        with open(txt_path, "w") as f:
            f.write(_generate_text_report(results))
        return txt_path

    doc = SimpleDocTemplate(filepath, pagesize=A4,
                            topMargin=20*mm, bottomMargin=20*mm,
                            leftMargin=20*mm, rightMargin=20*mm)
    styles = getSampleStyleSheet()

    # Custom styles
    title_style = ParagraphStyle('CustomTitle', parent=styles['Title'],
                                 textColor=HexColor('#00d4ff'), fontSize=24)
    heading_style = ParagraphStyle('CustomHeading', parent=styles['Heading1'],
                                   textColor=HexColor('#f1f5f9'), fontSize=16)
    subheading_style = ParagraphStyle('CustomSubheading', parent=styles['Heading2'],
                                      textColor=HexColor('#94a3b8'), fontSize=13)
    body_style = ParagraphStyle('CustomBody', parent=styles['Normal'],
                                textColor=HexColor('#333333'), fontSize=10)
    alert_style = ParagraphStyle('Alert', parent=styles['Normal'],
                                 textColor=HexColor('#ef4444'), fontSize=12, spaceAfter=6)

    story = []

    # Title page
    story.append(Spacer(1, 50*mm))
    story.append(Paragraph("🛡️ ThreatScope V2", title_style))
    story.append(Spacer(1, 5*mm))
    story.append(Paragraph("DFIR Analysis Report", heading_style))
    story.append(Spacer(1, 10*mm))
    story.append(Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", body_style))
    story.append(Paragraph(f"File: {results.get('metadata', {}).get('filepath', 'N/A')}", body_style))
    story.append(Paragraph(f"Analyzer: 0xSABRY ThreatScope V2", body_style))
    story.append(PageBreak())

    # Executive Summary
    story.append(Paragraph("Executive Summary", heading_style))
    story.append(Spacer(1, 5*mm))
    score = results.get("threat_score", 0)
    level = results.get("threat_level", "Unknown")
    summary = results.get("summary", {})
    story.append(Paragraph(f"<b>Threat Score:</b> {score}% — {level}", alert_style if score >= 60 else body_style))
    story.append(Paragraph(
        f"<b>Total Findings:</b> {summary.get('total_findings', 0)} | "
        f"<b>Critical:</b> {summary.get('critical', 0)} | "
        f"<b>High:</b> {summary.get('high', 0)} | "
        f"<b>Medium:</b> {summary.get('medium', 0)} | "
        f"<b>Low:</b> {summary.get('low', 0)}", body_style
    ))
    story.append(Paragraph(
        f"<b>MITRE Techniques:</b> {summary.get('mitre_techniques', 0)} | "
        f"<b>IOCs:</b> {summary.get('total_iocs', 0)} | "
        f"<b>Correlations:</b> {summary.get('correlations', 0)}", body_style
    ))
    story.append(Spacer(1, 10*mm))

    # Findings Table
    story.append(Paragraph("Critical and High Findings", heading_style))
    findings = results.get("findings", [])
    critical_high = [f for f in findings if f.get("severity", "").lower() in ("critical", "high")]

    if critical_high:
        table_data = [["Severity", "Title", "Description", "MITRE"]]
        for f in critical_high[:30]:
            table_data.append([
                f.get("severity", "").upper(),
                f.get("title", "")[:40],
                f.get("description", "")[:60],
                f.get("mitre", ""),
            ])

        t = Table(table_data, colWidths=[60, 120, 200, 70])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), HexColor('#1a2332')),
            ('TEXTCOLOR', (0, 0), (-1, 0), HexColor('#f1f5f9')),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, HexColor('#334155')),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [HexColor('#f8fafc'), HexColor('#e2e8f0')]),
        ]))
        story.append(t)

    # IOCs
    story.append(Spacer(1, 10*mm))
    story.append(Paragraph("Indicators of Compromise", heading_style))
    iocs = results.get("iocs", {})
    if iocs.get("top_ips"):
        story.append(Paragraph(f"<b>Top IPs:</b> {', '.join(iocs['top_ips'][:10])}", body_style))
    if iocs.get("top_domains"):
        story.append(Paragraph(f"<b>Top Domains:</b> {', '.join(iocs['top_domains'][:10])}", body_style))
    if iocs.get("cves"):
        story.append(Paragraph(f"<b>CVEs Found:</b> {', '.join(iocs['cves'][:10])}", body_style))

    doc.build(story)
    logger.info(f"PDF report exported to {filepath}")
    return filepath


def export_docx(results: dict, filepath: str) -> str:
    """Export analysis results as DOCX document."""
    if not DOCX_AVAILABLE:
        logger.warning("python-docx not installed — DOCX export unavailable")
        return export_json(results, filepath.replace(".docx", ".json"))

    doc = Document()

    # Title
    title = doc.add_heading("ThreatScope V2 — DFIR Analysis Report", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 212, 255)

    # Metadata
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"File: {results.get('metadata', {}).get('filepath', 'N/A')}")
    doc.add_paragraph(f"MD5: {results.get('metadata', {}).get('file_hashes', {}).get('md5', 'N/A')}")

    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    score = results.get("threat_score", 0)
    level = results.get("threat_level", "Unknown")
    summary = results.get("summary", {})
    doc.add_paragraph(f"Threat Score: {score}% ({level})")
    doc.add_paragraph(
        f"Findings: {summary.get('total_findings', 0)} total — "
        f"{summary.get('critical', 0)} critical, {summary.get('high', 0)} high, "
        f"{summary.get('medium', 0)} medium, {summary.get('low', 0)} low"
    )

    # Top Findings
    doc.add_heading("Key Findings", level=1)
    findings = results.get("findings", [])
    for f in findings[:20]:
        p = doc.add_paragraph()
        run = p.add_run(f"[{f.get('severity', '').upper()}] ")
        if f.get("severity", "").lower() == "critical":
            run.font.color.rgb = RGBColor(255, 23, 68)
        elif f.get("severity", "").lower() == "high":
            run.font.color.rgb = RGBColor(239, 68, 68)
        p.add_run(f"{f.get('title', '')} — {f.get('description', '')}")

    # IOCs
    doc.add_heading("Indicators of Compromise", level=1)
    iocs = results.get("iocs", {})
    if iocs.get("top_ips"):
        doc.add_paragraph(f"IPs: {', '.join(iocs['top_ips'][:10])}")
    if iocs.get("top_domains"):
        doc.add_paragraph(f"Domains: {', '.join(iocs['top_domains'][:10])}")

    doc.save(filepath)
    logger.info(f"DOCX report exported to {filepath}")
    return filepath


def _serialize_results(results: dict) -> dict:
    """Serialize results, converting sets to lists."""
    serialized = {}
    for key, value in results.items():
        if isinstance(value, set):
            serialized[key] = list(value)
        elif isinstance(value, dict):
            serialized[key] = _serialize_results(value)
        elif isinstance(value, list):
            serialized[key] = [
                _serialize_results(item) if isinstance(item, dict) else
                list(item) if isinstance(item, set) else item
                for item in value
            ]
        else:
            serialized[key] = value
    return serialized


def _generate_text_report(results: dict) -> str:
    """Generate a plain text report as fallback."""
    lines = [
        "=" * 70,
        "ThreatScope V2 — DFIR Analysis Report",
        "=" * 70,
        f"File: {results.get('metadata', {}).get('filepath', 'N/A')}",
        f"Threat Score: {results.get('threat_score', 0)}% ({results.get('threat_level', 'N/A')})",
        f"Total Findings: {results.get('summary', {}).get('total_findings', 0)}",
        "-" * 70,
    ]

    for f in results.get("findings", [])[:30]:
        lines.append(f"[{f.get('severity', '').upper()}] {f.get('title', '')} — {f.get('description', '')}")

    return "\n".join(lines)
